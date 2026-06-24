"""Real-behavior coverage for vincio.generation.templates.

Exercises type coercion branches, citation enforcement, required/default slot
resolution, strict vs. non-strict placeholder handling, and the DOCX merge-field
path (python-docx is installed). The PDF path is verified to require pypdf.
"""

from __future__ import annotations

import io
from pathlib import Path

import docx
import pytest

from vincio.core.errors import GenerationError
from vincio.generation.templates import (
    Slot,
    fill_docx_form,
    fill_pdf_form,
    fill_text_template,
)

# --------------------------------------------------------------------------- #
# Type coercion via rendered output
# --------------------------------------------------------------------------- #


def test_currency_coercion_formats_with_thousands_and_cents() -> None:
    out = fill_text_template(
        "{{amt}}",
        {"amt": 1234567.5},
        slots=[Slot(name="amt", type="currency")],
    )
    assert out == "$1,234,567.50"


def test_currency_coercion_non_numeric_falls_back_to_str() -> None:
    out = fill_text_template(
        "{{amt}}",
        {"amt": "free"},
        slots=[Slot(name="amt", type="currency")],
    )
    assert out == "free"


def test_number_integer_drops_decimal() -> None:
    out = fill_text_template(
        "{{n}}",
        {"n": 7.0},
        slots=[Slot(name="n", type="number")],
    )
    assert out == "7"


def test_number_non_integer_keeps_decimal() -> None:
    out = fill_text_template(
        "{{n}}",
        {"n": 3.5},
        slots=[Slot(name="n", type="number")],
    )
    assert out == "3.5"


def test_number_non_numeric_falls_back_to_str() -> None:
    # Hits the ValueError branch (lines 50-51).
    out = fill_text_template(
        "{{n}}",
        {"n": "twelve"},
        slots=[Slot(name="n", type="number")],
    )
    assert out == "twelve"


@pytest.mark.parametrize("truthy", [True, "true", "True", 1, "1", "yes"])
def test_bool_truthy_values_render_yes(truthy: object) -> None:
    out = fill_text_template(
        "{{flag}}",
        {"flag": truthy},
        slots=[Slot(name="flag", type="bool")],
    )
    assert out == "Yes"


@pytest.mark.parametrize("falsy", [False, "false", 0, "no", "anything"])
def test_bool_non_truthy_values_render_no(falsy: object) -> None:
    out = fill_text_template(
        "{{flag}}",
        {"flag": falsy},
        slots=[Slot(name="flag", type="bool", required=False)],
    )
    assert out == "No"


def test_none_value_coerces_to_empty_string() -> None:
    # Optional slot with no value and no default -> coerced to "" (line 40).
    out = fill_text_template(
        "[{{x}}]",
        {},
        slots=[Slot(name="x", required=False)],
    )
    assert out == "[]"


def test_string_type_passes_through_str() -> None:
    out = fill_text_template(
        "{{s}}",
        {"s": 42},
        slots=[Slot(name="s", type="string")],
    )
    assert out == "42"


# --------------------------------------------------------------------------- #
# Required / default resolution
# --------------------------------------------------------------------------- #


def test_required_slot_without_value_raises() -> None:
    with pytest.raises(GenerationError, match="required slot 'who' has no value"):
        fill_text_template(
            "Hi {{who}}",
            {},
            slots=[Slot(name="who", required=True)],
        )


def test_default_is_used_when_value_absent() -> None:
    out = fill_text_template(
        "Hi {{who}}",
        {},
        slots=[Slot(name="who", default="World")],
    )
    assert out == "Hi World"


def test_explicit_value_overrides_default() -> None:
    out = fill_text_template(
        "Hi {{who}}",
        {"who": "Ada"},
        slots=[Slot(name="who", default="World")],
    )
    assert out == "Hi Ada"


def test_ad_hoc_value_without_slot_renders() -> None:
    # No slot declared for the key -> Slot(name=key) default, not required-checked.
    out = fill_text_template("{{x}}", {"x": "v"})
    assert out == "v"


# --------------------------------------------------------------------------- #
# Citation enforcement
# --------------------------------------------------------------------------- #


def test_must_cite_without_citation_raises() -> None:
    with pytest.raises(
        GenerationError, match="slot 'claim' requires a citation but none was provided"
    ):
        fill_text_template(
            "{{claim}}",
            {"claim": "the sky is blue"},
            slots=[Slot(name="claim", must_cite=True)],
        )


def test_must_cite_with_valid_citation_passes() -> None:
    out = fill_text_template(
        "{{claim}}",
        {"claim": "Revenue grew [E1]"},
        slots=[Slot(name="claim", must_cite=True)],
        evidence_ids=["E1"],
    )
    assert out == "Revenue grew [E1]"


def test_must_cite_with_unknown_evidence_id_raises() -> None:
    # Hits the invalid-ids branch (lines 63->66).
    with pytest.raises(
        GenerationError, match=r"slot 'claim' cites unknown evidence: \['E9'\]"
    ):
        fill_text_template(
            "{{claim}}",
            {"claim": "Per the filing [E9]"},
            slots=[Slot(name="claim", must_cite=True)],
            evidence_ids=["E1"],
        )


def test_must_cite_without_evidence_id_whitelist_accepts_any_citation() -> None:
    # valid_ids is None -> the membership branch is skipped (63->exit).
    out = fill_text_template(
        "{{claim}}",
        {"claim": "Stated [E42]"},
        slots=[Slot(name="claim", must_cite=True)],
    )
    assert out == "Stated [E42]"


# --------------------------------------------------------------------------- #
# strict / non-strict placeholder handling
# --------------------------------------------------------------------------- #


def test_strict_unknown_placeholder_raises() -> None:
    with pytest.raises(
        GenerationError, match="template references unknown slot 'missing'"
    ):
        fill_text_template("Has {{missing}}", {})


def test_non_strict_leaves_unknown_placeholder_intact() -> None:
    out = fill_text_template("Has {{missing}}", {}, strict=False)
    assert out == "Has {{missing}}"


def test_whitespace_inside_placeholder_is_tolerated() -> None:
    out = fill_text_template("{{  name  }}", {"name": "x"})
    assert out == "x"


def test_multiple_placeholders_all_substituted() -> None:
    out = fill_text_template(
        "{{a}}-{{b}}-{{a}}",
        {"a": "1", "b": "2"},
    )
    assert out == "1-2-1"


# --------------------------------------------------------------------------- #
# DOCX merge-field path (python-docx installed)
# --------------------------------------------------------------------------- #


def _make_docx(tmp_path: Path, paragraph: str, cell: str | None = None) -> str:
    document = docx.Document()
    document.add_paragraph(paragraph)
    if cell is not None:
        table = document.add_table(rows=1, cols=1)
        table.rows[0].cells[0].paragraphs[0].add_run(cell)
    path = str(tmp_path / "template.docx")
    document.save(path)
    return path


def _docx_text(data: bytes) -> str:
    out = docx.Document(io.BytesIO(data))
    parts = [p.text for p in out.paragraphs]
    for table in out.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(p.text for p in cell.paragraphs)
    return "\n".join(parts)


def test_fill_docx_paragraph_and_table(tmp_path: Path) -> None:
    src = _make_docx(tmp_path, "Dear {{name}},", cell="Total {{amt}}")
    data = fill_docx_form(
        src,
        {"name": "Ada", "amt": 1000},
        slots=[Slot(name="amt", type="currency")],
    )
    text = _docx_text(data)
    assert "Dear Ada," in text
    assert "Total $1,000.00" in text
    assert "{{" not in text


def test_fill_docx_leaves_paragraph_without_placeholder_untouched(
    tmp_path: Path,
) -> None:
    src = _make_docx(tmp_path, "No placeholders here")
    data = fill_docx_form(src, {})
    assert "No placeholders here" in _docx_text(data)


def test_fill_docx_writes_output_path(tmp_path: Path) -> None:
    src = _make_docx(tmp_path, "Hi {{who}}")
    out_path = str(tmp_path / "out.docx")
    data = fill_docx_form(src, {"who": "Sam"}, output_path=out_path)
    on_disk = Path(out_path).read_bytes()
    assert on_disk == data
    assert "Hi Sam" in _docx_text(on_disk)


def test_fill_docx_unknown_placeholder_is_left_intact(tmp_path: Path) -> None:
    # The docx path's sub uses resolved.get(...) with the original as fallback,
    # so an undeclared placeholder survives rather than raising.
    src = _make_docx(tmp_path, "Keep {{ghost}}")
    data = fill_docx_form(src, {})
    assert "{{ghost}}" in _docx_text(data)


def test_fill_docx_propagates_required_slot_error(tmp_path: Path) -> None:
    src = _make_docx(tmp_path, "Hi {{who}}")
    with pytest.raises(GenerationError, match="required slot 'who' has no value"):
        fill_docx_form(src, {}, slots=[Slot(name="who", required=True)])


# --------------------------------------------------------------------------- #
# PDF path (pypdf not installed -> import-guard error)
# --------------------------------------------------------------------------- #


def test_fill_pdf_form_without_pypdf_raises_install_hint() -> None:
    try:
        import pypdf  # noqa: F401
    except ImportError:
        with pytest.raises(GenerationError, match="requires pypdf"):
            fill_pdf_form("nonexistent.pdf", {})
    else:  # pragma: no cover - pypdf present in some envs
        pytest.skip("pypdf installed; import-guard branch not exercised")
