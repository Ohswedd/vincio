"""Template & form filling with typed, citation-aware slots.

The same grounding guarantee the :class:`~vincio.generation.builder.DocumentBuilder`
applies to free-form documents, applied to fixed templates: a ``{{slot}}`` is
filled with a type-coerced value, and a slot marked ``must_cite`` must carry a
valid ``[E1]``-style citation or the fill fails. Text templates render
dependency-free; DOCX merge fields and PDF AcroForms fill behind the
``vincio[gen-docx]`` / ``vincio[pdf]`` extras.
"""

from __future__ import annotations

import re
from typing import Any, Literal

from pydantic import BaseModel

from ..core.errors import GenerationError
from ..output.parsers import extract_citations

__all__ = ["Slot", "fill_text_template", "fill_docx_form", "fill_pdf_form"]

SlotType = Literal["string", "number", "date", "currency", "bool"]

_PLACEHOLDER_RE = re.compile(r"\{\{\s*([\w.]+)\s*\}\}")


class Slot(BaseModel):
    """A typed, optionally citation-bearing template slot."""

    name: str
    type: SlotType = "string"
    required: bool = True
    must_cite: bool = False
    default: Any = None


def _coerce(value: Any, slot_type: SlotType) -> str:
    if value is None:
        return ""
    if slot_type == "currency":
        try:
            return f"${float(value):,.2f}"
        except (TypeError, ValueError):
            return str(value)
    if slot_type == "number":
        try:
            number = float(value)
            return str(int(number)) if number.is_integer() else str(number)
        except (TypeError, ValueError):
            return str(value)
    if slot_type == "bool":
        return "Yes" if value in (True, "true", "True", 1, "1", "yes") else "No"
    return str(value)


def _check_citations(
    name: str, rendered: str, valid_ids: set[str] | None
) -> None:
    citations = extract_citations(rendered)
    if not citations:
        raise GenerationError(f"slot {name!r} requires a citation but none was provided")
    if valid_ids is not None:
        invalid = [c for c in citations if c not in valid_ids]
        if invalid:
            raise GenerationError(f"slot {name!r} cites unknown evidence: {invalid}")


def _resolve_values(
    values: dict[str, Any],
    slots: list[Slot] | None,
    valid_ids: set[str] | None,
) -> dict[str, str]:
    slot_map = {s.name: s for s in (slots or [])}
    resolved: dict[str, str] = {}
    keys = set(values) | set(slot_map)
    for key in keys:
        slot = slot_map.get(key, Slot(name=key))
        raw = values.get(key, slot.default)
        if raw is None and slot.required and key in slot_map:
            raise GenerationError(f"required slot {key!r} has no value")
        rendered = _coerce(raw, slot.type)
        if slot.must_cite:
            # Check the raw (pre-coercion) value so a citation supplied alongside
            # a typed value is detected and the rendered value keeps its format.
            _check_citations(key, str(raw if raw is not None else ""), valid_ids)
        resolved[key] = rendered
    return resolved


def fill_text_template(
    template: str,
    values: dict[str, Any],
    *,
    slots: list[Slot] | None = None,
    evidence_ids: list[str] | None = None,
    strict: bool = True,
) -> str:
    """Fill ``{{slot}}`` placeholders in a text/Markdown/HTML template.

    Values are type-coerced per ``slots``; a ``must_cite`` slot must contain a
    valid citation. With ``strict`` (default), a placeholder with no value raises;
    otherwise it is left intact.
    """
    valid = set(evidence_ids) if evidence_ids is not None else None
    resolved = _resolve_values(values, slots, valid)

    def replace(match: re.Match[str]) -> str:
        key = match.group(1)
        if key in resolved:
            return resolved[key]
        if strict:
            raise GenerationError(f"template references unknown slot {key!r}")
        return match.group(0)

    return _PLACEHOLDER_RE.sub(replace, template)


def fill_docx_form(
    template_path: str,
    values: dict[str, Any],
    *,
    slots: list[Slot] | None = None,
    evidence_ids: list[str] | None = None,
    output_path: str | None = None,
) -> bytes:
    """Fill ``{{slot}}`` merge fields in a DOCX template, returning the bytes.

    Replaces placeholders in paragraphs and table cells. Writes to
    ``output_path`` when given.
    """
    try:
        import io

        import docx
    except ImportError as exc:  # pragma: no cover - optional dep
        raise GenerationError('DOCX form filling requires python-docx: pip install "vincio[gen-docx]"') from exc

    valid = set(evidence_ids) if evidence_ids is not None else None
    resolved = _resolve_values(values, slots, valid)

    def fill_paragraph(paragraph: Any) -> None:
        if "{{" not in paragraph.text:
            return
        filled = _PLACEHOLDER_RE.sub(lambda m: resolved.get(m.group(1), m.group(0)), paragraph.text)
        for run in list(paragraph.runs):
            run.text = ""
        if paragraph.runs:
            paragraph.runs[0].text = filled
        else:
            paragraph.add_run(filled)

    document = docx.Document(template_path)
    for paragraph in document.paragraphs:
        fill_paragraph(paragraph)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    fill_paragraph(paragraph)
    buffer = io.BytesIO()
    document.save(buffer)
    data = buffer.getvalue()
    if output_path:
        from pathlib import Path

        Path(output_path).write_bytes(data)
    return data


def fill_pdf_form(
    template_path: str,
    values: dict[str, Any],
    *,
    slots: list[Slot] | None = None,
    evidence_ids: list[str] | None = None,
    output_path: str | None = None,
) -> bytes:
    """Fill a PDF AcroForm's fields by name, returning the bytes.

    Uses pypdf's form-field update; writes to ``output_path`` when given.
    """
    try:
        import io

        from pypdf import PdfReader, PdfWriter
    except ImportError as exc:  # pragma: no cover - optional dep
        raise GenerationError('PDF form filling requires pypdf: pip install "vincio[pdf]"') from exc

    valid = set(evidence_ids) if evidence_ids is not None else None
    resolved = _resolve_values(values, slots, valid)

    reader = PdfReader(template_path)
    writer = PdfWriter()
    writer.append(reader)
    for page in writer.pages:
        writer.update_page_form_field_values(page, resolved)
    buffer = io.BytesIO()
    writer.write(buffer)
    data = buffer.getvalue()
    if output_path:
        from pathlib import Path

        Path(output_path).write_bytes(data)
    return data
