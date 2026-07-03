"""The document generation engine.

:class:`DocumentBuilder` turns a *validated* result — an
:class:`~vincio.output.schemas.OutputContract` output, a
:class:`~vincio.core.types.RunResult`, a structured mapping, Markdown, or a
:class:`~vincio.generation.model.DocumentModel` — into a rendered, structurally-
validated, provenance-stamped, audit-logged artifact. Because the input already
passed validation, the document is grounded by construction; the builder lays it
out, checks it against a :class:`~vincio.generation.contracts.DocumentContract`,
records a ``document_generate`` event carrying the source evidence ids, and never
invents content.
"""

from __future__ import annotations

import difflib
import re
from typing import TYPE_CHECKING, Any

from ..core.errors import DocumentContractError, GenerationError
from ..documents.parsers import TableData
from .contracts import DocumentContract, validate_document
from .model import DocBlock, DocumentModel
from .render import DocumentArtifact, RenderFormat, render

if TYPE_CHECKING:  # pragma: no cover - typing only
    from ..security.audit import AuditLog

__all__ = ["DocumentBuilder", "markdown_to_model", "generate_redline"]


_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_OL_RE = re.compile(r"^\s*\d+[.)]\s+(.*)$")
_UL_RE = re.compile(r"^\s*[-*+]\s+(.*)$")
_TABLE_ROW_RE = re.compile(r"^\s*\|(.+)\|\s*$")
_TABLE_SEP_RE = re.compile(r"^\s*\|?[\s:|-]+\|?\s*$")


def markdown_to_model(text: str, *, title: str = "") -> DocumentModel:
    """Parse Markdown into a :class:`DocumentModel`, preserving block order."""
    model = DocumentModel(title=title)
    lines = text.splitlines()
    index = 0
    paragraph: list[str] = []
    list_items: list[str] = []
    list_ordered = False

    def flush_paragraph() -> None:
        if paragraph:
            model.blocks.append(DocBlock(kind="paragraph", text=" ".join(paragraph).strip()))
            paragraph.clear()

    def flush_list() -> None:
        nonlocal list_ordered
        if list_items:
            model.blocks.append(DocBlock(kind="list", items=list(list_items), ordered=list_ordered))
            list_items.clear()
            list_ordered = False

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        heading = _HEADING_RE.match(line)
        # Code fence
        if stripped.startswith("```"):
            flush_paragraph()
            flush_list()
            language = stripped[3:].strip()
            body: list[str] = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                body.append(lines[index])
                index += 1
            model.blocks.append(DocBlock(kind="code", text="\n".join(body), language=language))
            index += 1
            continue
        # Table
        row_match = _TABLE_ROW_RE.match(line)
        if row_match and index + 1 < len(lines) and _TABLE_SEP_RE.match(lines[index + 1]):
            flush_paragraph()
            flush_list()
            cols = [c.strip() for c in row_match.group(1).split("|")]
            rows: list[list[str]] = []
            cursor = index + 2
            while cursor < len(lines):
                next_match = _TABLE_ROW_RE.match(lines[cursor])
                if not next_match:
                    break
                rows.append([c.strip() for c in next_match.group(1).split("|")])
                cursor += 1
            model.blocks.append(DocBlock(kind="table", table=TableData(columns=cols, rows=rows)))
            index = cursor
            continue
        if heading:
            flush_paragraph()
            flush_list()
            level = len(heading.group(1))
            head_text = heading.group(2).strip()
            # A leading H1 with no model title becomes the document title.
            if level == 1 and not model.title and not model.blocks:
                model.title = head_text
            else:
                model.blocks.append(DocBlock(kind="heading", text=head_text, level=max(1, level - 1)))
            index += 1
            continue
        ol = _OL_RE.match(line)
        ul = _UL_RE.match(line)
        list_match = ol or ul
        if list_match is not None:
            flush_paragraph()
            if list_items and list_ordered != bool(ol):
                flush_list()
            list_ordered = bool(ol)
            list_items.append(list_match.group(1).strip())
            index += 1
            continue
        if stripped.startswith(">"):
            flush_paragraph()
            flush_list()
            model.blocks.append(DocBlock(kind="quote", text=stripped.lstrip(">").strip()))
            index += 1
            continue
        if not stripped:
            flush_paragraph()
            flush_list()
            index += 1
            continue
        flush_list()
        paragraph.append(stripped)
        index += 1
    flush_paragraph()
    flush_list()
    return model


def _table_from_value(value: Any) -> TableData | None:
    """Coerce a value into a TableData (dict with columns/rows, or list of dicts)."""
    if isinstance(value, TableData):
        return value
    if isinstance(value, dict) and ("columns" in value or "rows" in value):
        return TableData(
            id=str(value.get("id", "")),
            title=str(value.get("title", "")),
            columns=[str(c) for c in value.get("columns", [])],
            rows=[[str(c) for c in row] for row in value.get("rows", [])],
        )
    if isinstance(value, list) and value and all(isinstance(r, dict) for r in value):
        columns = list(dict.fromkeys(k for row in value for k in row))
        rows = [[str(row.get(c, "")) for c in columns] for row in value]
        return TableData(columns=columns, rows=rows)
    return None


def mapping_to_model(data: dict[str, Any], *, title: str = "") -> DocumentModel:
    """Build a model from a structured mapping.

    Recognizes ``title``/``subtitle`` and a ``sections`` list whose entries carry
    ``heading``/``title``, ``level``, ``body``/``text``/``paragraphs``,
    ``items``, and ``table``. A mapping without ``sections`` is rendered as a
    titled key→value section so any validated dict yields a usable document.
    """
    model = DocumentModel(title=title or str(data.get("title", "")), subtitle=str(data.get("subtitle", "")))
    sections = data.get("sections")
    if isinstance(sections, list):
        for entry in sections:
            if not isinstance(entry, dict):
                model.blocks.append(DocBlock(kind="paragraph", text=str(entry)))
                continue
            head = str(entry.get("heading") or entry.get("title") or "").strip()
            if head:
                model.blocks.append(
                    DocBlock(kind="heading", text=head, level=int(entry.get("level", 1)),
                             anchor=str(entry.get("anchor", "")))
                )
            body = entry.get("body") or entry.get("text") or ""
            paragraphs = entry.get("paragraphs")
            if isinstance(paragraphs, list):
                for para in paragraphs:
                    model.blocks.append(DocBlock(kind="paragraph", text=str(para)))
            elif body:
                model.blocks.append(DocBlock(kind="paragraph", text=str(body)))
            items = entry.get("items") or entry.get("bullets")
            if isinstance(items, list) and items:
                model.blocks.append(DocBlock(kind="list", items=[str(i) for i in items]))
            table = _table_from_value(entry.get("table"))
            if table is not None:
                model.blocks.append(DocBlock(kind="table", table=table))
        return model
    # No explicit sections: render remaining keys as a definition section.
    for key, value in data.items():
        if key in ("title", "subtitle"):
            continue
        model.blocks.append(DocBlock(kind="heading", text=str(key).replace("_", " ").title(), level=1))
        table = _table_from_value(value)
        if table is not None:
            model.blocks.append(DocBlock(kind="table", table=table))
        elif isinstance(value, list):
            model.blocks.append(DocBlock(kind="list", items=[str(v) for v in value]))
        else:
            model.blocks.append(DocBlock(kind="paragraph", text=str(value)))
    return model


class DocumentBuilder:
    """Render validated results into cited, contract-checked, audited documents."""

    def __init__(
        self,
        *,
        audit_log: AuditLog | None = None,
        tenant_id: str | None = None,
    ) -> None:
        self.audit_log = audit_log
        self.tenant_id = tenant_id

    # -- IR construction -------------------------------------------------------

    def to_model(self, source: Any, *, title: str = "") -> DocumentModel:
        """Coerce a validated result into a :class:`DocumentModel`."""
        if isinstance(source, DocumentModel):
            if title and not source.title:
                source.title = title
            return source
        # RunResult (duck-typed to avoid an import cycle): prefer structured
        # output, fall back to raw text; carry the evidence ids for provenance.
        if hasattr(source, "raw_text") and hasattr(source, "evidence"):
            evidence_ids = [getattr(e, "id", "") for e in getattr(source, "evidence", [])]
            # Truthy (not just non-None) so an empty structured container falls
            # back to the raw text instead of rendering an empty document.
            payload = source.output if getattr(source, "output", None) else source.raw_text
            model = self.to_model(payload, title=title)
            model.source_evidence_ids = [e for e in evidence_ids if e]
            return model
        if hasattr(source, "model_dump"):
            return mapping_to_model(source.model_dump(mode="python"), title=title)
        if isinstance(source, dict):
            return mapping_to_model(source, title=title)
        if isinstance(source, list):
            return mapping_to_model({"sections": source}, title=title)
        if isinstance(source, str):
            return markdown_to_model(source, title=title)
        raise GenerationError(f"cannot build a document from {type(source).__name__}")

    # -- build -----------------------------------------------------------------

    def build(
        self,
        source: Any,
        *,
        format: RenderFormat = "markdown",
        contract: DocumentContract | None = None,
        title: str = "",
        evidence_ids: list[str] | None = None,
        footnotes: list[str] | None = None,
        bibliography: list[str] | None = None,
        run_id: str | None = None,
    ) -> DocumentArtifact:
        """Build, validate, render, and audit a document in one call."""
        model = self.to_model(source, title=title)
        if footnotes:
            model.footnotes = list(footnotes)
        if bibliography:
            model.bibliography = list(bibliography)
        if evidence_ids:
            model.source_evidence_ids = list(dict.fromkeys(model.source_evidence_ids + evidence_ids))

        report = None
        if contract is not None:
            report = validate_document(model, contract)
            if not report.valid:
                raise DocumentContractError(
                    "document does not satisfy its contract: " + "; ".join(report.violations),
                    violations=report.violations,
                )

        artifact = render(model, format)
        if report is not None:
            artifact.metadata["contract_repairs"] = report.repairs
        self._audit(artifact, model)
        return artifact

    def _audit(self, artifact: DocumentArtifact, model: DocumentModel) -> None:
        if self.audit_log is None:
            return
        self.audit_log.record(
            "document_generate",
            tenant_id=self.tenant_id,
            resource=artifact.title or "document",
            details={
                "format": artifact.format,
                "media_type": artifact.media_type,
                "bytes": len(artifact.content),
                # frozen audit-detail key — external consumers bind to it.
                "content_sha256": artifact.digest(),
                "source_evidence_ids": list(model.source_evidence_ids),
                "blocks": len(model.blocks),
                "words": model.word_count(),
            },
        )


# -- redline ------------------------------------------------------------------


def _tokenize(text: str) -> list[str]:
    return re.findall(r"\S+\s*", text)


def _redline_ops(original: str, revised: str) -> list[tuple[str, str]]:
    """Word-level diff as a list of (op, text) where op is equal/insert/delete."""
    a = _tokenize(original)
    b = _tokenize(revised)
    matcher = difflib.SequenceMatcher(a=a, b=b, autojunk=False)
    ops: list[tuple[str, str]] = []
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            ops.append(("equal", "".join(a[i1:i2])))
        else:
            if tag in ("delete", "replace"):
                ops.append(("delete", "".join(a[i1:i2])))
            if tag in ("insert", "replace"):
                ops.append(("insert", "".join(b[j1:j2])))
    return ops


def generate_redline(
    original: str,
    revised: str,
    *,
    format: RenderFormat = "docx",
    title: str = "Redline",
) -> DocumentArtifact:
    """Generate a tracked-change redline between two texts.

    Pairs the ``DOCUMENT_COMPARISON`` intent with a marked-up artifact:
    DOCX/PDF carry visual tracked changes (insertions underlined, deletions
    struck through); Markdown/HTML use ``**ins**`` / ``~~del~~`` conventions.
    """
    ops = _redline_ops(original, revised)
    if format == "docx":
        return _redline_docx(ops, title)
    # Text formats: build a model with an inline-marked paragraph.
    marked: list[str] = []
    for op, text in ops:
        if not text:
            continue
        if op == "equal":
            marked.append(text)
        elif op == "insert":
            marked.append(f"**{text.strip()}** ")
        else:
            marked.append(f"~~{text.strip()}~~ ")
    model = DocumentModel(title=title)
    model.paragraph("".join(marked).strip())
    return render(model, format)


def _redline_docx(ops: list[tuple[str, str]], title: str) -> DocumentArtifact:
    try:
        import io

        import docx
        from docx.shared import RGBColor
    except ImportError as exc:  # pragma: no cover - optional dep
        raise GenerationError('DOCX redline requires python-docx: pip install "vincio[gen-docx]"') from exc

    document = docx.Document()
    document.add_heading(title, level=0)
    paragraph = document.add_paragraph()
    for op, text in ops:
        if not text:
            continue
        run = paragraph.add_run(text)
        if op == "insert":
            run.font.underline = True
            run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)  # green
        elif op == "delete":
            run.font.strike = True
            run.font.color.rgb = RGBColor(0xB7, 0x1C, 0x1C)  # red
    buffer = io.BytesIO()
    document.save(buffer)
    from .render import MEDIA_TYPES

    return DocumentArtifact(
        format="docx", media_type=MEDIA_TYPES["docx"], content=buffer.getvalue(), title=title
    )
