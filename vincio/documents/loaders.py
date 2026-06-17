"""Document loaders.

``load_document(path)`` dispatches through the :mod:`~vincio.documents.registry`
parser registry and returns a structured :class:`Document` with text, sections,
tables, and provenance. Built-in formats: text, Markdown, a real-parser HTML
path (with table extraction), CSV/TSV, structured JSON/JSONL/YAML, code,
email (.eml/mbox/.msg), PDF (``vincio[pdf]``; OCR auto-fallback for scanned
pages), DOCX (``vincio[docx]``), XLSX (openpyxl), PPTX/EPUB/RTF/ODT
(dependency-free), Parquet (``vincio[parquet]``), and code repositories.
``load_media(path, transcriber=...)`` ingests audio as a timestamped transcript.
"""

from __future__ import annotations

import inspect
import json
import re
from email import policy as email_policy
from email.parser import BytesParser
from pathlib import Path
from typing import Any

from ..core.errors import DocumentError, LoaderError
from ..core.types import Document, EvidenceItem
from ..core.utils import new_id
from . import formats as _formats  # noqa: F401 — registers PPTX/EPUB/RTF/ODT/Parquet/mbox/.msg
from .parsers import (
    extract_code_symbols,
    extract_markdown_sections,
    extract_markdown_tables,
    parse_csv_table,
    parse_html,
    strip_html,
    structure_data,
)
from .registry import default_parser_registry, register_loader

__all__ = [
    "load_document",
    "load_directory",
    "load_media",
    "load_pdf",
    "load_docx",
    "load_xlsx",
    "figure_evidence",
    "register_loader",
    "SUPPORTED_EXTENSIONS",
    "supported_extensions",
]

_CODE_EXTENSIONS = {
    ".py": "python",
    ".js": "javascript",
    ".ts": "typescript",
    ".tsx": "typescript",
    ".jsx": "javascript",
    ".go": "go",
    ".rs": "rust",
    ".java": "java",
    ".rb": "ruby",
    ".c": "c",
    ".h": "c",
    ".cpp": "cpp",
    ".cs": "csharp",
    ".sql": "sql",
    ".sh": "shell",
}


def _read_text(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return path.read_text(encoding="latin-1")


def _resolve_maybe_async(value: Any) -> Any:
    """Run an awaitable to completion from sync code; pass values through."""
    if inspect.isawaitable(value):
        from ..providers.base import run_sync

        return run_sync(value)
    return value


# -- registered built-in loaders ----------------------------------------------


@register_loader(".html", ".htm")
def _load_html(path: Path, **_: Any) -> Document:
    html = _read_text(path)
    title, text, sections, tables = parse_html(html)
    return Document(
        text=text,
        title=title or path.stem,
        media_type="text/html",
        sections=[s.model_dump(mode="json") for s in sections],
        tables=[t.model_dump(mode="json") for t in tables],
    )


@register_loader(".csv", ".tsv")
def _load_csv(path: Path, **_: Any) -> Document:
    content = _read_text(path)
    delimiter = "\t" if path.suffix.lower() == ".tsv" else None
    table = parse_csv_table(content, title=path.stem, delimiter=delimiter)
    return Document(
        text=table.to_text(),
        title=path.stem,
        media_type="text/csv",
        tables=[table.model_dump(mode="json")],
    )


@register_loader(".json", ".jsonl")
def _load_json(path: Path, **_: Any) -> Document:
    content = _read_text(path)
    suffix = path.suffix.lower()
    parsed: Any
    try:
        if suffix == ".jsonl":
            parsed = [json.loads(line) for line in content.splitlines() if line.strip()]
        else:
            parsed = json.loads(content)
    except json.JSONDecodeError:
        return Document(
            text=content, title=path.stem, media_type="application/json",
            metadata={"parsed": False},
        )
    text, sections, tables = structure_data(parsed, title=path.stem)
    return Document(
        text=text or content,
        title=path.stem,
        media_type="application/json",
        sections=[s.model_dump(mode="json") for s in sections],
        tables=[t.model_dump(mode="json") for t in tables],
        metadata={"parsed": True},
    )


@register_loader(".yaml", ".yml")
def _load_yaml(path: Path, **_: Any) -> Document:
    content = _read_text(path)
    try:
        import yaml

        parsed = yaml.safe_load(content)
    except Exception:  # noqa: BLE001 - any YAML error falls back to raw text
        return Document(text=content, title=path.stem, media_type="application/yaml")
    text, sections, tables = structure_data(parsed, title=path.stem)
    return Document(
        text=text or content,
        title=path.stem,
        media_type="application/yaml",
        sections=[s.model_dump(mode="json") for s in sections],
        tables=[t.model_dump(mode="json") for t in tables],
    )


@register_loader(*_CODE_EXTENSIONS)
def _load_code(path: Path, **_: Any) -> Document:
    source = _read_text(path)
    language = _CODE_EXTENSIONS[path.suffix.lower()]
    symbols = extract_code_symbols(source, language=language)
    return Document(
        text=source,
        title=path.name,
        media_type=f"text/x-{language}",
        metadata={
            "language": language,
            "symbols": [s.model_dump(mode="json") for s in symbols],
            "imports": [s.name for s in symbols if s.kind == "import"],
        },
    )


@register_loader(".eml")
def _load_email(path: Path, **_: Any) -> Document:
    message = BytesParser(policy=email_policy.default).parse(path.open("rb"))
    body = message.get_body(preferencelist=("plain", "html"))
    text = ""
    if body is not None:
        content = body.get_content()
        text = strip_html(content) if body.get_content_subtype() == "html" else content
    headers = {
        "from": str(message.get("From", "")),
        "to": str(message.get("To", "")),
        "date": str(message.get("Date", "")),
        "subject": str(message.get("Subject", "")),
    }
    return Document(
        text=f"Subject: {headers['subject']}\nFrom: {headers['from']}\nTo: {headers['to']}\n\n{text}",
        title=headers["subject"] or path.stem,
        media_type="message/rfc822",
        metadata=headers,
    )


@register_loader(".txt", ".md", ".rst", ".markdown")
def _load_text(path: Path, **_: Any) -> Document:
    suffix = path.suffix.lower()
    text = _read_text(path)
    document = Document(
        text=text,
        title=path.stem,
        media_type="text/markdown" if suffix in (".md", ".markdown", ".rst") else "text/plain",
    )
    if suffix in (".md", ".markdown", ".rst"):
        document.sections = [s.model_dump(mode="json") for s in extract_markdown_sections(text)]
        document.tables = [t.model_dump(mode="json") for t in extract_markdown_tables(text)]
    return document


# -- core dispatch ------------------------------------------------------------


def supported_extensions() -> set[str]:
    """The live set of loadable suffixes — reflects loaders registered at any
    time via :func:`register_loader`, plus the built-in PDF/DOCX/XLSX paths.
    Prefer this over the import-time :data:`SUPPORTED_EXTENSIONS` snapshot."""
    return default_parser_registry().suffixes() | {".pdf", ".docx", ".xlsx"}


# Import-time snapshot of the built-in suffixes (kept for backward compatibility);
# use supported_extensions() to see loaders registered after import.
SUPPORTED_EXTENSIONS = supported_extensions()


def load_document(
    path: str | Path,
    *,
    tenant_id: str | None = None,
    layout: bool = False,
    ocr_engine: Any = None,
) -> Document:
    """Load a single file into a Document. Raises LoaderError on failure.

    Dispatches through the parser registry. ``layout=True`` selects the
    layout-aware PDF path; ``ocr_engine`` routes low-text-yield PDF pages through
    OCR (see :func:`load_pdf`). The dependency-free text path is the default for
    every format.
    """
    path = Path(path)
    if not path.is_file():
        raise LoaderError(f"file not found: {path}")
    suffix = path.suffix.lower()
    registry = default_parser_registry()
    try:
        if suffix == ".pdf":
            document = load_pdf(path, layout=layout, ocr_engine=ocr_engine)
        elif suffix == ".docx":
            document = load_docx(path)
        elif suffix == ".xlsx":
            document = load_xlsx(path)
        elif registry.supports(suffix):
            document = registry.load(path, layout=layout, ocr_engine=ocr_engine)
        else:  # unknown suffix → best-effort plain text
            document = _load_text(path)
    except LoaderError:
        raise
    except Exception as exc:  # noqa: BLE001 - wrap any parser failure
        raise LoaderError(f"failed to load {path}: {exc}") from exc

    document.source_uri = str(path)
    document.tenant_id = tenant_id
    document.metadata.setdefault("filename", path.name)
    document.metadata.setdefault("size_bytes", path.stat().st_size)
    return document


def _default_rasterizer(path: str | Path, page_index: int) -> str:
    """Rasterize one PDF page to a temp PNG via pypdfium2; return its path."""
    try:
        import pypdfium2  # type: ignore
    except ImportError as exc:  # pragma: no cover - optional dep
        raise LoaderError(
            'PDF OCR fallback requires a rasterizer: pip install "vincio[ocr]" '
            "(pypdfium2), or pass a custom rasterizer"
        ) from exc
    import tempfile

    pdf = pypdfium2.PdfDocument(str(path))
    try:
        page = pdf[page_index]
        bitmap = page.render(scale=2.0)
        image = bitmap.to_pil()
        handle = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
        image.save(handle.name)
        return handle.name
    finally:
        pdf.close()


def load_pdf(
    path: str | Path,
    *,
    layout: bool = False,
    ocr_engine: Any = None,
    ocr_min_chars: int = 16,
    rasterizer: Any = None,
) -> Document:
    """PDF text extraction via pypdf (``pip install "vincio[pdf]"``).

    With ``layout=True``, uses the layout-aware path (reading order, tables,
    figures) via ``vincio[pdf-layout]``. With an ``ocr_engine`` (a
    :class:`~vincio.documents.ocr.OCREngine`), pages whose extracted text is
    below ``ocr_min_chars`` are rasterized (``rasterizer`` or pypdfium2) and
    OCR'd, with ``extractor='ocr'`` recorded per page so provenance stays honest.
    """
    if layout:
        from .layout import extract_pdf_layout

        return extract_pdf_layout(path)
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise LoaderError('PDF support requires pypdf: pip install "vincio[pdf]"') from exc
    path = Path(path)
    reader = PdfReader(str(path))
    rasterize = rasterizer or _default_rasterizer
    pages: list[str] = []
    sections: list[dict[str, Any]] = []
    ocr_pages: list[int] = []
    for page_number, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        extractor = "text"
        if ocr_engine is not None and len(page_text.strip()) < ocr_min_chars:
            try:
                image_path = rasterize(path, page_number - 1)
                ocr_text = _resolve_maybe_async(ocr_engine.extract_text(image_path)) or ""
                if len(ocr_text.strip()) > len(page_text.strip()):
                    page_text = ocr_text
                    extractor = "ocr"
                    ocr_pages.append(page_number)
            except (LoaderError, DocumentError):
                raise
            except Exception:  # noqa: BLE001 - OCR is best-effort per page
                pass
        pages.append(page_text)
        sections.append(
            {"title": f"page {page_number}", "level": 1, "path": [f"p{page_number}"],
             "text": page_text, "start_line": 0, "page": page_number, "extractor": extractor}
        )
    metadata: dict[str, Any] = {"page_count": len(reader.pages)}
    if ocr_pages:
        metadata["ocr_pages"] = ocr_pages
    if reader.metadata:
        for key in ("title", "author", "subject"):
            value = getattr(reader.metadata, key, None)
            if value:
                metadata[key] = str(value)
    return Document(
        text="\n\n".join(pages),
        title=metadata.get("title") or path.stem,
        media_type="application/pdf",
        sections=sections,
        metadata=metadata,
    )


def load_docx(path: str | Path) -> Document:
    """DOCX extraction via python-docx (``pip install "vincio[docx]"``)."""
    try:
        import docx  # python-docx
    except ImportError as exc:
        raise LoaderError('DOCX support requires python-docx: pip install "vincio[docx]"') from exc
    path = Path(path)
    document = docx.Document(str(path))
    paragraphs: list[str] = []
    sections: list[dict[str, Any]] = []
    current: dict[str, Any] | None = None
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style = (paragraph.style.name or "").lower() if paragraph.style else ""
        if style.startswith("heading"):
            level = int(re.sub(r"\D", "", style) or 1)
            current = {"title": text, "level": level, "path": [text], "text": "", "start_line": 0}
            sections.append(current)
        elif current is not None:
            current["text"] = (current["text"] + "\n" + text).strip()
        paragraphs.append(text)
    tables: list[dict[str, Any]] = []
    for index, table in enumerate(document.tables, start=1):
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if rows:
            tables.append(
                {"id": f"T{index}", "title": "", "columns": rows[0], "rows": rows[1:],
                 "source": str(path), "footnotes": [], "units": {}, "inferred_schema": {}, "quality": {}}
            )
    return Document(
        text="\n".join(paragraphs),
        title=path.stem,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        sections=sections,
        tables=tables,
    )


def load_xlsx(path: str | Path) -> Document:
    """XLSX extraction via openpyxl, including formulas and sheet names."""
    try:
        import openpyxl
    except ImportError as exc:
        raise LoaderError("XLSX support requires openpyxl: pip install openpyxl") from exc
    path = Path(path)
    workbook = openpyxl.load_workbook(str(path), data_only=False)
    tables: list[dict[str, Any]] = []
    text_parts: list[str] = []
    for sheet_index, sheet in enumerate(workbook.worksheets, start=1):
        rows = [[("" if c.value is None else str(c.value)) for c in row] for row in sheet.iter_rows()]
        rows = [row for row in rows if any(cell.strip() for cell in row)]
        if not rows:
            continue
        from .parsers import TableData, infer_table_schema, table_quality_checks

        table = TableData(
            id=f"S{sheet_index}",
            title=sheet.title,
            columns=rows[0],
            rows=rows[1:],
            source=f"{path.name}:{sheet.title}",
        )
        table.inferred_schema = infer_table_schema(table)
        table.quality = table_quality_checks(table)
        formulas = [
            f"{cell.coordinate}={cell.value}"
            for row in sheet.iter_rows()
            for cell in row
            if isinstance(cell.value, str) and cell.value.startswith("=")
        ]
        table_dump = table.model_dump(mode="json")
        table_dump["formulas"] = formulas
        tables.append(table_dump)
        text_parts.append(table.to_text())
    return Document(
        text="\n\n".join(text_parts),
        title=path.stem,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        tables=tables,
        metadata={"sheets": [s.title for s in workbook.worksheets]},
    )


# -- audio / media ------------------------------------------------------------


def load_media(
    path: str | Path,
    *,
    transcriber: Any,
    tenant_id: str | None = None,
) -> Document:
    """Ingest audio/video as a timestamped transcript :class:`Document`.

    ``transcriber`` is a :class:`~vincio.documents.audio.Transcriber` (e.g.
    ``MockTranscriber`` offline, ``WhisperTranscriber`` online). Each transcript
    segment becomes a section carrying its timestamp and (when diarized) speaker.
    """
    path = Path(path)
    if not path.is_file():
        raise LoaderError(f"file not found: {path}")
    transcript = _resolve_maybe_async(transcriber.transcribe(path))
    sections: list[dict[str, Any]] = []
    for index, segment in enumerate(transcript.segments, start=1):
        prefix = segment.timestamp + (f" {segment.speaker}:" if segment.speaker else "")
        sections.append(
            {"title": f"segment {index}", "level": 1, "path": [f"seg{index}"],
             "text": f"{prefix} {segment.text}".strip(), "start_line": 0,
             "start": segment.start, "end": segment.end, "speaker": segment.speaker,
             "extractor": "transcript"}
        )
    return Document(
        source_uri=str(path),
        text=transcript.text,
        title=path.stem,
        media_type="audio/transcript",
        sections=sections,
        tenant_id=tenant_id,
        metadata={
            "extractor": "transcript",
            "language": transcript.language,
            "duration_s": transcript.duration_s,
            "segment_count": len(sections),
            "filename": path.name,
        },
    )


# -- figures → evidence -------------------------------------------------------


def figure_evidence(
    document: Document,
    *,
    crops: dict[int, str],
    analyzer: Any = None,
    ocr_engine: Any = None,
) -> list[EvidenceItem]:
    """Turn cropped PDF figure regions into citable evidence with bounding boxes.

    ``crops`` maps a figure index (into ``document.metadata['figures']``) to a
    cropped image path. Each crop is described by an
    :class:`~vincio.documents.multimodal.ImageAnalyzer` (``analyzer``) or
    transcribed by an :class:`~vincio.documents.ocr.OCREngine` (``ocr_engine``),
    yielding evidence stamped with the figure's bbox and page.
    """
    figures = document.metadata.get("figures") or []
    source_id = document.id
    items: list[EvidenceItem] = []
    for fig_index, image_path in crops.items():
        figure = figures[fig_index] if 0 <= fig_index < len(figures) else {}
        bbox = figure.get("bbox")
        page = figure.get("page")
        if analyzer is not None:
            observations = _resolve_maybe_async(analyzer.observe(image_path))
            for obs_index, obs in enumerate(observations, start=1):
                items.append(
                    EvidenceItem(
                        id=f"{source_id}:FIG{fig_index}:R{obs_index}",
                        source_id=source_id,
                        source_type="image",
                        text=f"[figure {fig_index}] {obs.observation}",
                        media_ref=image_path,
                        page=page,
                        authority=obs.confidence,
                        provenance=0.9,
                        metadata={"bbox": bbox, "region": obs.region, "figure_index": fig_index},
                    )
                )
        elif ocr_engine is not None:
            text = _resolve_maybe_async(ocr_engine.extract_text(image_path)) or ""
            if text.strip():
                items.append(
                    EvidenceItem(
                        id=f"{source_id}:FIG{fig_index}",
                        source_id=source_id,
                        source_type="image",
                        text=f"[figure {fig_index}] {text.strip()}",
                        media_ref=image_path,
                        page=page,
                        provenance=0.9,
                        metadata={"bbox": bbox, "figure_index": fig_index, "extractor": "ocr"},
                    )
                )
    return items


_DEFAULT_IGNORES = {
    ".git", ".hg", ".svn", "node_modules", "__pycache__", ".venv", "venv",
    "dist", "build", ".mypy_cache", ".pytest_cache", ".ruff_cache",
}


def load_directory(
    path: str | Path,
    *,
    extensions: set[str] | None = None,
    ignore_dirs: set[str] | None = None,
    max_files: int = 5000,
    tenant_id: str | None = None,
) -> list[Document]:
    """Load a directory tree (code repositories / doc folders).

    Adds a repository summary document with the dependency/import graph when
    code files are present.
    """
    root = Path(path)
    if not root.is_dir():
        raise LoaderError(f"directory not found: {root}")
    ignore = _DEFAULT_IGNORES | (ignore_dirs or set())
    # Live registry, so loaders registered after import are honored here too.
    allowed = extensions or supported_extensions()
    documents: list[Document] = []
    import_graph: dict[str, list[str]] = {}
    for file_path in sorted(root.rglob("*")):
        if len(documents) >= max_files:
            break
        if not file_path.is_file():
            continue
        if any(part in ignore for part in file_path.parts):
            continue
        if file_path.suffix.lower() not in allowed:
            continue
        try:
            document = load_document(file_path, tenant_id=tenant_id)
        except LoaderError:
            continue  # unreadable/optional-dep file inside a bulk load
        document.metadata["relative_path"] = str(file_path.relative_to(root))
        documents.append(document)
        if "imports" in document.metadata:
            import_graph[document.metadata["relative_path"]] = document.metadata["imports"]
    if import_graph:
        summary_lines = ["Repository import graph:"]
        for module, imports in sorted(import_graph.items()):
            if imports:
                summary_lines.append(f"{module} -> {', '.join(sorted(set(imports)))}")
        documents.append(
            Document(
                id=new_id("doc"),
                text="\n".join(summary_lines),
                title="repository_summary",
                media_type="text/plain",
                source_uri=str(root),
                tenant_id=tenant_id,
                metadata={"kind": "repo_summary", "file_count": len(documents)},
            )
        )
    return documents
